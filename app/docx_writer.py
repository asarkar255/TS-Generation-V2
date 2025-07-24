
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def add_page_break(doc):
    doc.add_page_break()

def add_heading(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 255)
    run.font.size = Pt(14)
    paragraph.space_after = Pt(12)

def add_subheading(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text.strip(":"))
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 255)
    run.font.size = Pt(12)
    paragraph.space_after = Pt(6)

def add_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.space_after = Pt(6)
    cursor = 0
    bold_matches = list(re.finditer(r"\*\*(.+?)\*\*", text))
    if not bold_matches:
        paragraph.add_run(text)
    else:
        for match in bold_matches:
            start, end = match.span()
            paragraph.add_run(text[cursor:start])
            bold_run = paragraph.add_run(match.group(1))
            bold_run.bold = True
            cursor = end
        paragraph.add_run(text[cursor:])

def add_code_block(doc, code_lines):
    para = doc.add_paragraph()
    run = para.add_run("\n".join(code_lines))
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    para.space_after = Pt(6)

def add_markdown_table(doc, lines):
    # Normalize and split all rows
    rows = [row.strip().strip("|").split("|") for row in lines]
    rows = [[cell.strip() for cell in row] for row in rows if row]

    if len(rows) < 1:
        return

    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = 'Table Grid'
    for i, header in enumerate(rows[0]):
        table.cell(0, i).text = header

    for row in rows[1:]:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = cell

def create_docx(ts_text: str, buffer):
    doc = Document()
    doc.add_heading('TECHNICAL SPECIFICATION', level=1)

    lines = ts_text.splitlines()
    current_section = ""
    current_content = []
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []

    def flush_current_content():
        if current_section:
            add_heading(doc, current_section)
        for para in current_content:
            if para.strip().startswith("**") and para.strip().endswith(":**"):
                add_subheading(doc, para.strip(" *:"))
            elif para.strip().startswith("**") and para.strip().endswith("**"):
                add_subheading(doc, para.strip(" *"))
            else:
                add_paragraph(doc, para)

    section_header_pattern = re.compile(r"^\d+\.\s+[A-Z ]+$")
    inline_subheader_pattern = re.compile(r"^\*\*(.+?):\*\*$")
    page_break_marker = re.compile(r"^PAGE \d+", re.IGNORECASE)

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Page Break
        if page_break_marker.match(line):
            flush_current_content()
            current_content = []
            add_page_break(doc)
            continue

        # Code Block
        if line.startswith("```"):
            in_code_block = not in_code_block
            if not in_code_block:
                add_code_block(doc, code_block_lines)
                code_block_lines = []
            continue
        elif in_code_block:
            code_block_lines.append(line)
            continue

        # Table Block
        if line.startswith("|") and line.endswith("|"):
            table_lines.append(line)
            in_table = True
            continue
        elif in_table and not line.startswith("|"):
            flush_current_content()
            current_content = []
            add_markdown_table(doc, table_lines)
            table_lines = []
            in_table = False

        # Section Header
        if section_header_pattern.match(line):
            flush_current_content()
            current_section = line
            current_content = []
            continue

        current_content.append(line)

    # Final flush
    if current_section and current_content:
        flush_current_content()

    doc.save(buffer)
